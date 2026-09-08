from dataclasses import replace
from datetime import date
from io import BytesIO
import unittest

from docx import Document

from backend.services.manual_dispatch.delivery_docket_docx_parser import (
    apply_delivery_docket_validation,
    extract_delivery_docket_docx_text,
    parse_delivery_docket_docx_bytes,
    parse_delivery_docket_text,
)
from backend.services.manual_dispatch.delivery_suburb_region_service import (
    UNKNOWN_DELIVERY_AREA_WARNING,
)


IMPORT_DATE = date(2026, 8, 13)
CASE_4409_PARAGRAPHS = [
    "DELIVERY DOCKET: 4409/186066",
    "DATED: 03/09/2026",
    "VIA SYMES TRANSPORT 5443 4199 RING TO ADVISE",
    "DROP OFF TO SUNSHINE DEPOT BELOW",
    "E: miCK@SYMESTRANSPORT.COM",
    "DELIVERY TO:",
    "SYMES TSPT C/- ECO WOOL INSULATION",
    "BUILDING 3, 82 -86 BERKSHIRE RD SUNSHINE NTH",
    "ENTER VIA STEERS ST THRU GATE ON RHS",
    "ON FORWARD TO:    **PLEASE CHARGE CUSTOMER ACCT#",
    "K2 INDUSTRIAL SUPPLIES",
    "55 CORNELIA CREEK RD",
    "ECHUCA  5482 4446",
    "ORDER NUMBER: 30010822",
    "63 X 10KG COLOUR RAGS",
    "1 PALLET",
    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS PH; 03 9930 7700",
]


def _docx_bytes(lines, table_rows=()):
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    if table_rows:
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        for values in table_rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, values):
                cell.text = value
    payload = BytesIO()
    document.save(payload)
    return payload.getvalue()


def _parse(lines, filename):
    return parse_delivery_docket_docx_bytes(
        _docx_bytes(lines),
        source_filename=filename,
        import_date=IMPORT_DATE,
    )


class DeliveryDocketDocxParserTest(unittest.TestCase):
    def test_extracts_paragraphs_and_table_cells_in_document_order(self):
        payload = _docx_bytes(
            ["DELIVERY DOCKET: 4370/185467", "DATED: 10/08/2026"],
            table_rows=[("DELIVER TO:", "PATONS TRANSPORT"), ("143 FOUNDATION ROAD", "TRUGANINA")],
        )

        text = extract_delivery_docket_docx_text(payload)

        self.assertLess(text.index("DATED: 10/08/2026"), text.index("DELIVER TO:"))
        self.assertLess(text.index("DELIVER TO:"), text.index("PATONS TRANSPORT"))
        self.assertLess(text.index("143 FOUNDATION ROAD"), text.index("TRUGANINA"))

    def test_case_4370_prefers_substantive_header_and_on_forward_customer(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 073",
                "DELIVERY DOCKET: 4370/185467",
                "DATED: 10/08/2026",
                "EMAIL PATONS TRANSPORT TO ADVISE DROP OFF OF STOCK",
                "admin@patonstransport.com.au",
                "DELIVER to",
                "PATONS TRANSPORT",
                "c/o Victorian freight specialists",
                "143 foundation road",
                "truganina",
                "ON FWD TO:",
                "CHEMBLAST INDUSTRIAL COATINGS",
                "14 DEANE STREET",
                "BUNBURY WA 6230",
                "52 x 10kg COLOUR SINGLET RAGS",
                "1 PALLET",
                "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS",
                "PH: 03 9930 7700",
            ],
            "docket-4370.docx",
        )

        self.assertEqual("4370", parsed.docket_number)
        self.assertEqual("185467", parsed.docket_reference)
        self.assertEqual("185467", parsed.invoice_number)
        self.assertEqual("2026-08-10", parsed.invoice_date)
        self.assertIsNone(parsed.order_no)
        self.assertEqual("CHEMBLAST INDUSTRIAL COATINGS", parsed.company_name)
        self.assertIsNone(parsed.phone)
        self.assertEqual("143 FOUNDATION ROAD", parsed.delivery_address)
        self.assertEqual("TRUGANINA", parsed.suburb)
        self.assertIsNone(parsed.postcode)
        self.assertEqual("2026-08-14", parsed.delivery_date)
        self.assertEqual((1, 0, 0), (parsed.pallet_quantity, parsed.loose_bags_quantity, parsed.carton_quantity))
        self.assertEqual(
            {
                "product_code": None,
                "product_name": "COLOUR SINGLET RAGS",
                "quantity": 520,
                "unit": "KG",
                "package_quantity": 52,
                "package_unit": "BAG10",
            },
            parsed.product_lines[0],
        )
        self.assertIn("Delivery Docket: 4370", parsed.note)
        self.assertIn("Deliver To: PATONS TRANSPORT", parsed.note)
        self.assertIn("C/O VICTORIAN FREIGHT SPECIALISTS", parsed.note)
        self.assertIn("On Forward To: CHEMBLAST INDUSTRIAL COATINGS", parsed.note)
        self.assertIn("Booking Email: admin@patonstransport.com.au", parsed.note)
        self.assertNotIn("03 9930 7700", parsed.note)
        self.assertNotIn("Delivery Docket: 073", parsed.note)

    def test_case_4371_keeps_newway_reference_without_guessing_invoice(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 4371/NEWWAY 182-2",
                "Princes SS",
                "DATED: 10/08/2026",
                "EMAIL SCT LOGISTICS TO ADVISE DROP OFF OF STOCK TO THEIR ALTONA DEPOT",
                "melbourne.pickups2@sctlogistics.com.au",
                "DELIVER to",
                "SCT ALTONA",
                "7 WESTLINK COURT",
                "ALTONA VIC 3018",
                "SCT REFERENCE# 103010642",
                "ON FWD TO CUSTOMER:",
                "PRINCES FABRICARE MAROOCHYDORE",
                "BAUHINA CENTRE, 3/526 MAROOCHYDORE ROAD",
                "KUNDA PARK",
                "QLD 4556",
                "PHONE: 07 5445 5133",
                "TOTAL: 2 PALLETS",
            ],
            "docket-4371.docx",
        )

        self.assertEqual("4371", parsed.docket_number)
        self.assertEqual("NEWWAY 182-2 Princes SS", parsed.docket_reference)
        self.assertIsNone(parsed.invoice_number)
        self.assertEqual("2026-08-10", parsed.invoice_date)
        self.assertEqual("PRINCES FABRICARE MAROOCHYDORE", parsed.company_name)
        self.assertEqual("07 5445 5133", parsed.phone)
        self.assertEqual("7 WESTLINK COURT", parsed.delivery_address)
        self.assertEqual("ALTONA", parsed.suburb)
        self.assertEqual("3018", parsed.postcode)
        self.assertEqual(2, parsed.pallet_quantity)
        self.assertEqual([], parsed.product_lines)
        self.assertTrue(any("No product lines" in warning for warning in parsed.warnings))
        self.assertTrue(parsed.importable)
        self.assertIn("Docket Reference: NEWWAY 182-2 Princes SS", parsed.note)
        self.assertIn("SCT Reference: 103010642", parsed.note)
        self.assertIn("Booking Email: melbourne.pickups2@sctlogistics.com.au", parsed.note)

    def test_case_4372_ignores_template_header_and_keeps_attention_context(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 073",
                "DELIVERY DOCKET: 4372/NEWWAY 181-3, 185, 186 Princes AD",
                "DATED: 10/08/2026",
                "DELIVER to",
                "SCT ALTONA",
                "7 WESTLINK COURT",
                "ALTONA VIC 3018",
                "SCT REFERENCE# 103010643",
                "ON FWD TO CUSTOMER:",
                "PRINCES LINEN ADELAIDE",
                "ATTN: MELANIE FISHER",
                "15 WALDAREE STREET",
                "GEPPS CROSS SA 5094",
                "PHONE: 08 8121 5488",
                "TOTAL: 10 PALLETS",
            ],
            "docket-4372.docx",
        )

        self.assertEqual("4372", parsed.docket_number)
        self.assertIsNone(parsed.invoice_number)
        self.assertEqual("PRINCES LINEN ADELAIDE", parsed.company_name)
        self.assertEqual("08 8121 5488", parsed.phone)
        self.assertEqual("7 WESTLINK COURT", parsed.delivery_address)
        self.assertEqual("ALTONA", parsed.suburb)
        self.assertEqual("3018", parsed.postcode)
        self.assertEqual(10, parsed.pallet_quantity)
        self.assertIn("Contact: MELANIE FISHER", parsed.note)
        self.assertIn("GEPPS CROSS SA 5094", parsed.note)

    def test_case_4373_splits_final_customer_from_physical_depot(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 073",
                "DELIVERY DOCKET: 4373/185504",
                "DATED: 11/08/2026",
                "DELIVER to",
                "C/-LUBRIMAXX SUNSHINE",
                "30 SPENCER STREET",
                "SUNSHINE WEST",
                "ON FWD TO:",
                "NOEL'S AUTO PARTS",
                "366 EDWARD STREET",
                "WAGGA WAGGA NSW 2650",
                "ATT: DAVID",
                "PH: 02 6925 3777",
                "ORDER NUMBER: 40592",
                "36x10kgs COLOUR T-SHIRT RAGS",
                "1 X PALLET",
            ],
            "docket-4373.docx",
        )

        self.assertEqual("185504", parsed.invoice_number)
        self.assertEqual("40592", parsed.order_no)
        self.assertEqual("NOEL'S AUTO PARTS", parsed.company_name)
        self.assertEqual("02 6925 3777", parsed.phone)
        self.assertEqual("30 SPENCER STREET", parsed.delivery_address)
        self.assertEqual("SUNSHINE WEST", parsed.suburb)
        self.assertEqual(360, parsed.product_lines[0]["quantity"])
        self.assertIn("Deliver To: C/-LUBRIMAXX SUNSHINE", parsed.note)
        self.assertIn("On Forward Address: 366 EDWARD STREET / WAGGA WAGGA NSW 2650", parsed.note)
        self.assertIn("Contact: DAVID", parsed.note)

    def test_case_4374_parses_open_window_and_two_products_without_loose_double_count(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 073",
                "DELIVERY DOCKET: 4374/185503",
                "Email to book:",
                "office@keatingfreightlines.com.au",
                "DATED: 11/08/2026",
                "DELIVER to",
                "keatings transport OPEN 7AM-3pm",
                "36-38 glenbarry road",
                "campbellfield",
                "ON FWD TO:",
                "FASTENER SPECIALISTS",
                "274 TOWNSEND ST",
                "SOUTH ALBURY 2640",
                "STOCK:",
                "63 X 10KG WORKSHOP rags",
                "63 X 10KG WINDCHEATER RAGS",
                "2 X PALLETS",
            ],
            "docket-4374.docx",
        )

        self.assertEqual("FASTENER SPECIALISTS", parsed.company_name)
        self.assertIsNone(parsed.phone)
        self.assertEqual("36-38 GLENBARRY ROAD", parsed.delivery_address)
        self.assertEqual("CAMPBELLFIELD", parsed.suburb)
        self.assertEqual(("07:00", "15:00"), (parsed.start_time, parsed.end_time))
        self.assertEqual(2, parsed.pallet_quantity)
        self.assertEqual(0, parsed.loose_bags_quantity)
        self.assertEqual([630, 630], [line["quantity"] for line in parsed.product_lines])
        self.assertEqual(["BAG10", "BAG10"], [line["package_unit"] for line in parsed.product_lines])
        self.assertIn("Delivery Window: OPEN 7AM-3PM", parsed.note)
        self.assertIn("Booking Email: office@keatingfreightlines.com.au", parsed.note)

    def test_case_4375_preserves_order_slash_and_entry_as_note_only(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 4375/185512",
                "DATED: 11/08/2026",
                "DELIVER TO:",
                "JJS WASTE & RECYLING",
                "46-52 ELLIOTT ROAD",
                "ENTRY VIA 427 HAMMOND RD",
                "DANDENONG SOUTH 3175",
                "ORDER NUMBER: 77058/VIC",
                "45X10KG COLOURED SINGLET",
                "1pallet",
                "INVOICE TO FOLLOW FROM SMITHS RAGS",
            ],
            "docket-4375.docx",
        )

        self.assertEqual("DIRECT", parsed.delivery_mode)
        self.assertEqual("77058/VIC", parsed.order_no)
        self.assertEqual("JJS WASTE & RECYLING", parsed.company_name)
        self.assertEqual("46-52 ELLIOTT ROAD", parsed.delivery_address)
        self.assertEqual("DANDENONG SOUTH", parsed.suburb)
        self.assertEqual("3175", parsed.postcode)
        self.assertEqual(450, parsed.product_lines[0]["quantity"])
        self.assertIn("Entry: VIA 427 HAMMOND RD", parsed.note)
        self.assertNotIn("Deliver To Address: ENTRY", parsed.note)

    def test_case_4376_explicit_time_slot_wins_and_supplier_phone_is_excluded(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 4376/185531",
                "DATED: 21/08/2026",
                "DELIVER TO:",
                "SRGS PTY LTD",
                "MELBOURNE DISTRIBUTION CENTRE",
                "STORE DC31",
                "413 MT ATKINSON ROAD",
                "TRUGANINA",
                "Time slot: FRIDAY 21/08/26 @ 9am",
                "BOOKING # 2232671",
                "ORDER NUMBER: 4522009179",
                "1120 X 1.5KG COLOUR RAGS",
                "45 X 10KG COLOR T SHIRT RAGS",
                "6 PALLETS",
                "INVOICE TO FOLLOW FROM MCC RAGMAN PTY LTD",
                "TRADING AS MELBOURNE CLEANING CLOTHS",
                "PH: 03 9930 7700",
            ],
            "docket-4376.docx",
        )

        self.assertEqual("DIRECT", parsed.delivery_mode)
        self.assertEqual("2026-08-21", parsed.invoice_date)
        self.assertEqual("2026-08-14", parsed.delivery_date)
        self.assertEqual("09:00", parsed.start_time)
        self.assertIsNone(parsed.end_time)
        self.assertEqual("SRGS PTY LTD", parsed.company_name)
        self.assertIsNone(parsed.phone)
        self.assertEqual("413 MT ATKINSON ROAD", parsed.delivery_address)
        self.assertEqual("TRUGANINA", parsed.suburb)
        self.assertEqual(6, parsed.pallet_quantity)
        self.assertEqual([1680, 450], [line["quantity"] for line in parsed.product_lines])
        self.assertEqual(["BAG1.5", "BAG10"], [line["package_unit"] for line in parsed.product_lines])
        self.assertIn("Site: MELBOURNE DISTRIBUTION CENTRE", parsed.note)
        self.assertIn("Store: DC31", parsed.note)
        self.assertIn("Booking #: 2232671", parsed.note)
        self.assertIn("Time Slot: FRIDAY 21/08/26 @ 9AM", parsed.note)
        self.assertNotIn("03 9930 7700", parsed.note)

    def test_case_4391_supports_drop_off_on_forward_phone_and_bare_window(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 4391/185816",
                "DATED: 25/08/2026",
                "VIA BURKINSHAWS TRANSPORT",
                "DROP OFF TO:",
                "C/-CHEMHAUL LOGISTICS",
                "13 SOMERS STREET",
                "SUNSHINE NORTH",
                "8AM-3PM",
                "ON FORWARD to:",
                "UNITED FASTENERS",
                "UNIT 2/3 BALL PL",
                "WAGGA WAGGA PH: 02 9131 3333",
                "ORDER NUMBER: 40041845",
                "63 X10kgs WORKSHOP RAGS",
                "1 pallet",
            ],
            "docket-4391.docx",
        )

        self.assertEqual("4391", parsed.docket_number)
        self.assertEqual("185816", parsed.invoice_number)
        self.assertEqual("40041845", parsed.order_no)
        self.assertEqual("UNITED FASTENERS", parsed.company_name)
        self.assertEqual("02 9131 3333", parsed.phone)
        self.assertEqual("13 SOMERS STREET", parsed.delivery_address)
        self.assertEqual("SUNSHINE NORTH", parsed.suburb)
        self.assertIsNone(parsed.postcode)
        self.assertEqual(("08:00", "15:00"), (parsed.start_time, parsed.end_time))
        self.assertEqual(1, parsed.pallet_quantity)
        self.assertEqual(630, parsed.product_lines[0]["quantity"])
        self.assertTrue(parsed.importable)

    def test_case_4392_isolates_inline_phone_from_state_only_suburb(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 4392/185817",
                "DATED: 25/08/2026",
                "DELIVER to",
                "ABLETTS TRANSPORT (TRANSPORT DEPOT)",
                "47 ENDEAVOUR WAY",
                "SUNSHINE VIC ph: 9313 9933",
                "ON FWD TO:",
                "CAPITAL PAINT PLACE",
                "18 ISA STREET",
                "FYSHWICK ACT 2609",
                "44 X 10KG WHITE T SHIRT RAGS",
                "19 X 10KG FLANNEL RAGS",
                "1 PALLET",
            ],
            "docket-4392.docx",
        )

        self.assertEqual("4392", parsed.docket_number)
        self.assertEqual("185817", parsed.invoice_number)
        self.assertEqual("CAPITAL PAINT PLACE", parsed.company_name)
        self.assertEqual("9313 9933", parsed.phone)
        self.assertEqual("47 ENDEAVOUR WAY", parsed.delivery_address)
        self.assertEqual("SUNSHINE", parsed.suburb)
        self.assertIsNone(parsed.postcode)
        self.assertNotEqual("9933", parsed.postcode)
        self.assertNotEqual("18 ISA STREET", parsed.delivery_address)
        self.assertEqual([440, 190], [line["quantity"] for line in parsed.product_lines])
        self.assertTrue(parsed.importable)

    def test_case_4393_strips_address_annotation_and_parses_open_start_only(self):
        parsed = _parse(
            [
                "DELIVERY DOCKET: 4393/18511",
                "DATED: 25/08/2026",
                "DELIVER to",
                "SOUTH WEST FREIGHT (SWF) OPEN 7AM",
                "C/-CDM TRANSPORT",
                "47 ENDEAVOUR WAY **please charge smiths rags acct**",
                "SUNSHINE WEST VIC",
                "ON FWD TO CUSTOMER:",
                "Konnect fastening systems",
                "321 commercial st west",
                "Mount gambier sa 5290",
                "Caine m: 0428 723 989",
                "ORDER NUMBER:",
                "63 x10kg workshop/mix cotton",
                "1PALLET",
            ],
            "docket-4393.docx",
        )

        self.assertEqual("4393", parsed.docket_number)
        self.assertEqual("18511", parsed.docket_reference)
        self.assertIsNone(parsed.invoice_number)
        self.assertEqual("KONNECT FASTENING SYSTEMS", parsed.company_name)
        self.assertEqual("0428 723 989", parsed.phone)
        self.assertEqual("47 ENDEAVOUR WAY", parsed.delivery_address)
        self.assertNotIn("PLEASE CHARGE", parsed.delivery_address)
        self.assertEqual("SUNSHINE WEST", parsed.suburb)
        self.assertIsNone(parsed.postcode)
        self.assertEqual("07:00", parsed.start_time)
        self.assertIsNone(parsed.end_time)
        self.assertEqual(630, parsed.product_lines[0]["quantity"])
        self.assertTrue(parsed.importable)

    def test_case_4409_parses_on_forward_long_form_inline_suburb_and_phone(self):
        parsed = _parse(CASE_4409_PARAGRAPHS, "docket-4409.docx")

        self.assertEqual("4409", parsed.docket_number)
        self.assertEqual("186066", parsed.docket_reference)
        self.assertEqual("186066", parsed.invoice_number)
        self.assertEqual("2026-09-03", parsed.invoice_date)
        self.assertEqual("ON_FORWARD", parsed.delivery_mode)
        self.assertEqual("K2 INDUSTRIAL SUPPLIES", parsed.company_name)
        self.assertEqual("BUILDING 3, 82 -86 BERKSHIRE RD", parsed.delivery_address)
        self.assertEqual("SUNSHINE NTH", parsed.suburb)
        self.assertIsNone(parsed.postcode)
        self.assertEqual("5482 4446", parsed.phone)
        self.assertEqual("30010822", parsed.order_no)
        self.assertEqual(1, parsed.pallet_quantity)
        self.assertEqual(630, parsed.product_lines[0]["quantity"])
        self.assertEqual("COLOUR RAGS", parsed.product_lines[0]["product_name"])
        self.assertEqual("KG", parsed.product_lines[0]["unit"])
        self.assertEqual(63, parsed.product_lines[0]["package_quantity"])
        self.assertEqual("BAG10", parsed.product_lines[0]["package_unit"])
        self.assertTrue(parsed.importable)
        for context in (
            "Deliver To: SYMES TSPT C/- ECO WOOL INSULATION",
            "BUILDING 3, 82 -86 BERKSHIRE RD",
            "ENTER VIA STEERS ST THRU GATE ON RHS",
            "On Forward To: K2 INDUSTRIAL SUPPLIES",
            "On Forward Address: 55 CORNELIA CREEK RD / ECHUCA",
            "**PLEASE CHARGE CUSTOMER ACCT#",
            "VIA SYMES TRANSPORT 5443 4199 RING TO ADVISE",
            "Booking Email: mick@symestransport.com",
        ):
            self.assertIn(context, parsed.note)
        for invalid_profile_text in (
            "ON FORWARD TO", "PLEASE CHARGE CUSTOMER ACCT",
            "ENTER VIA STEERS ST THRU GATE ON RHS",
        ):
            self.assertNotIn(invalid_profile_text, parsed.company_name)
            self.assertNotIn(invalid_profile_text, parsed.suburb)
        self.assertNotIn("03 9930 7700", parsed.note)

    def test_on_forward_heading_variants_share_section_boundaries(self):
        for heading in (
            "ON FWD TO", "ON FWD TO CUSTOMER",
            "ON FORWARD TO", "ON FORWARD TO CUSTOMER",
        ):
            for annotation in ("", " **PLEASE CHARGE CUSTOMER ACCT#"):
                with self.subTest(heading=heading, annotation=annotation):
                    lines = list(CASE_4409_PARAGRAPHS)
                    lines[0] = "DELIVERY DOCKET: 4990/199990"
                    lines[5] = "DELIVER TO:"
                    lines[9] = f"{heading.lower()}:{annotation}"
                    parsed = _parse(lines, "heading-variants.docx")
                    self.assertEqual("ON_FORWARD", parsed.delivery_mode)
                    self.assertEqual("K2 INDUSTRIAL SUPPLIES", parsed.company_name)
                    self.assertEqual("BUILDING 3, 82 -86 BERKSHIRE RD", parsed.delivery_address)
                    self.assertEqual("SUNSHINE NTH", parsed.suburb)
                    self.assertEqual("5482 4446", parsed.phone)

                    reference = parse_delivery_docket_text(
                        f"DELIVERY DOCKET: 4990/NEWWAY 182-2\n{lines[9]}\nFINAL CUSTOMER",
                        import_date=IMPORT_DATE,
                    )
                    self.assertEqual("NEWWAY 182-2", reference.docket_reference)

    def test_inline_street_suburb_split_preserves_separate_address_formats(self):
        cases = (
            (["143 FOUNDATION ROAD", "TRUGANINA"], "143 FOUNDATION ROAD", "TRUGANINA", None),
            (["30 SPENCER STREET", "SUNSHINE WEST"], "30 SPENCER STREET", "SUNSHINE WEST", None),
            (["36-38 GLENBARRY ROAD", "CAMPBELLFIELD"], "36-38 GLENBARRY ROAD", "CAMPBELLFIELD", None),
            (["UNIT 2, 8-10 INDUSTRIAL AVE COBURG NORTH"], "UNIT 2, 8-10 INDUSTRIAL AVE", "COBURG NORTH", None),
            (["12 SUPPLY DRIVE AUBURN NSW 2144"], "12 SUPPLY DRIVE", "AUBURN", "2144"),
            (["321 COMMERCIAL ST WEST", "MOUNT GAMBIER SA 5290"], "321 COMMERCIAL ST WEST", "MOUNT GAMBIER", "5290"),
        )
        for address_lines, address, suburb, postcode in cases:
            with self.subTest(address_lines=address_lines):
                parsed = _parse([
                    "DELIVERY DOCKET: 4991/199991", "DELIVERY TO:", "TEST DEPOT",
                    *address_lines, "1 PALLET",
                ], "street-suburb.docx")
                self.assertEqual(address, parsed.delivery_address)
                self.assertEqual(suburb, parsed.suburb)
                self.assertEqual(postcode, parsed.postcode)

    def test_inline_phone_and_postcode_formats_are_unambiguous(self):
        cases = (
            ("ECHUCA 5482 4446", "ECHUCA", None, "5482 4446"),
            ("ECHUCA 03 5482 4446", "ECHUCA", None, "03 5482 4446"),
            ("WAGGA WAGGA 02 6925 3777", "WAGGA WAGGA", None, "02 6925 3777"),
            ("BUNBURY WA 6230", "BUNBURY", "6230", None),
            ("ALTONA VIC 3018", "ALTONA", "3018", None),
            ("GEPPS CROSS SA 5094", "GEPPS CROSS", "5094", None),
            ("ECHUCA 3564", "ECHUCA", "3564", None),
            ("ECHUCA PHONE: 03 5482 4446", "ECHUCA", None, "03 5482 4446"),
            ("ECHUCA PH: 5482 4446", "ECHUCA", None, "5482 4446"),
            ("ECHUCA TEL: 03 5482 4446", "ECHUCA", None, "03 5482 4446"),
            ("ECHUCA MOBILE: 0428 723 989", "ECHUCA", None, "0428 723 989"),
        )
        for line, suburb, postcode, phone in cases:
            with self.subTest(line=line):
                parsed = _parse([
                    "DELIVERY DOCKET: 4992/199992", "DELIVER TO:", "TEST CUSTOMER",
                    "55 CORNELIA CREEK RD", line, "1 PALLET",
                    "INVOICE TO FOLLOW FROM MCC", "PH: 03 9930 7700",
                ], "inline-phone.docx")
                self.assertEqual(suburb, parsed.suburb)
                self.assertEqual(postcode, parsed.postcode)
                self.assertEqual(phone, parsed.phone)

    def test_instruction_lines_are_context_not_profile_fields(self):
        for instruction in (
            "ENTER VIA STEERS ST THRU GATE ON RHS", "ENTRY VIA 427 HAMMOND RD",
            "RING TO ADVISE", "DROP OFF TO DEPOT BELOW", "PLEASE CHARGE CUSTOMER ACCT#",
        ):
            with self.subTest(instruction=instruction):
                parsed = _parse([
                    "DELIVERY DOCKET: 4993/199993", "DELIVER TO:",
                    instruction, "TEST CUSTOMER", instruction,
                    "143 FOUNDATION ROAD", instruction, "TRUGANINA", "1 PALLET",
                ], "instructions.docx")
                self.assertEqual("TEST CUSTOMER", parsed.company_name)
                self.assertEqual("143 FOUNDATION ROAD", parsed.delivery_address)
                self.assertEqual("TRUGANINA", parsed.suburb)
                self.assertIsNone(parsed.postcode)
                expected_note = instruction.replace("ENTRY ", "Entry: ", 1)
                self.assertIn(expected_note, parsed.note)

    def test_inline_address_does_not_guess_suburb_from_instructions_or_direction(self):
        for address_line in (
            "30 EXAMPLE ROAD WEST", "30 EXAMPLE ROAD RING TO ADVISE",
            "THE ROAD TO THE DEPOT",
        ):
            with self.subTest(address_line=address_line):
                parsed = _parse([
                    "DELIVERY DOCKET: 4994/199994", "DELIVER TO:", "TEST CUSTOMER",
                    address_line, "1 PALLET",
                ], "no-guessed-suburb.docx")
                self.assertIsNone(parsed.suburb)
                self.assertFalse(parsed.importable)

    def test_time_slot_date_does_not_override_import_business_date(self):
        parsed = parse_delivery_docket_text(
            """
            DELIVERY DOCKET: 4998/199998
            DATED: 20/08/2026
            DELIVER TO:
            BUSINESS DATE CUSTOMER
            1 EXAMPLE ROAD
            RICHMOND 3121
            TIME SLOT: MONDAY 10/08/26 @ 9AM
            5 X 10KG SAMPLE RAGS
            1 PALLET
            """,
            source_filename="business-date.docx",
            import_date=date(2026, 8, 21),
        )

        self.assertEqual("2026-08-24", parsed.delivery_date)
        self.assertEqual("09:00", parsed.start_time)
        self.assertIn("Time Slot: MONDAY 10/08/26 @ 9AM", parsed.note)

    def test_fractional_actual_kg_is_warned_and_not_importable(self):
        parsed = parse_delivery_docket_text(
            """
            DELIVERY DOCKET: 4999/199999
            DATED: 12/08/2026
            DELIVER TO:
            EXAMPLE CUSTOMER
            1 EXAMPLE ROAD
            RICHMOND 3121
            3 X 1.5KG SAMPLE RAGS
            """,
            source_filename="fractional.docx",
            import_date=IMPORT_DATE,
        )

        self.assertFalse(parsed.importable)
        self.assertTrue(any("fractional" in warning.lower() for warning in parsed.warnings))

    def test_manual_corrections_revalidate_required_fields_without_clearing_review_warnings(self):
        parsed = parse_delivery_docket_text(
            """
            DELIVERY DOCKET: 4391/185816
            DATED: 25/08/2026
            DELIVER TO:
            UNITED FASTENERS
            1 TEST ROAD
            SUNSHINE NORTH
            ORDER NUMBER: 40041845
            1 PALLET
            """,
            source_filename="docket-4391.docx",
            import_date=date(2026, 8, 25),
        )
        unrelated_warning = "Informational docket warning."
        repaired = replace(
            parsed,
            company_name=None,
            delivery_address=None,
            suburb=None,
            warnings=[
                "Customer company was not found.",
                "Deliver To street address was not found.",
                "Deliver To suburb was not found.",
                unrelated_warning,
                UNKNOWN_DELIVERY_AREA_WARNING,
            ],
            importable=False,
            selected=False,
        )

        repaired.company_name = "UNITED FASTENERS"
        apply_delivery_docket_validation(repaired)
        self.assertFalse(repaired.importable)
        self.assertNotIn("Customer company was not found.", repaired.warnings)
        self.assertIn("Deliver To street address was not found.", repaired.warnings)
        self.assertIn("Deliver To suburb was not found.", repaired.warnings)

        repaired.suburb = "SUNSHINE NORTH"
        apply_delivery_docket_validation(repaired)
        self.assertFalse(repaired.importable)
        self.assertNotIn("Deliver To suburb was not found.", repaired.warnings)

        repaired.delivery_address = "1 TEST ROAD"
        apply_delivery_docket_validation(repaired)
        self.assertTrue(repaired.importable)
        self.assertFalse(repaired.selected)
        self.assertEqual(
            [unrelated_warning, UNKNOWN_DELIVERY_AREA_WARNING],
            repaired.warnings,
        )

        repaired.selected = True
        repaired.suburb = ""
        apply_delivery_docket_validation(repaired)
        self.assertFalse(repaired.importable)
        self.assertFalse(repaired.selected)
        self.assertIn("Deliver To suburb was not found.", repaired.warnings)

    def test_revalidation_preserves_load_duplicate_and_product_blockers(self):
        parsed = parse_delivery_docket_text(
            """
            DELIVERY DOCKET: 4391/185816
            DATED: 25/08/2026
            DELIVER TO:
            UNITED FASTENERS
            1 TEST ROAD
            SUNSHINE NORTH
            1 PALLET
            """,
            source_filename="docket-4391.docx",
            import_date=date(2026, 8, 25),
        )
        parsed.warnings.append(UNKNOWN_DELIVERY_AREA_WARNING)
        apply_delivery_docket_validation(parsed)
        self.assertTrue(parsed.importable)

        no_load = replace(
            parsed,
            pallet_quantity=0,
            loose_bags_quantity=0,
            carton_quantity=0,
            warnings=list(parsed.warnings),
        )
        apply_delivery_docket_validation(no_load)
        self.assertFalse(no_load.importable)
        self.assertIn("No pallet, loose bag, or carton load was found.", no_load.warnings)

        duplicate = replace(
            parsed,
            is_duplicate=True,
            warnings=[*parsed.warnings, "Duplicate invoice number already exists."],
        )
        apply_delivery_docket_validation(duplicate)
        self.assertFalse(duplicate.importable)
        self.assertFalse(duplicate.selected)
        self.assertIn("Duplicate invoice number already exists.", duplicate.warnings)

        fractional = replace(
            parsed,
            product_lines=[{
                "product_name": "SAMPLE RAGS",
                "quantity": 4.5,
                "unit": "KG",
                "package_quantity": 3,
                "package_unit": "BAG1.5",
            }],
            warnings=list(parsed.warnings),
        )
        apply_delivery_docket_validation(fractional)
        self.assertFalse(fractional.importable)
        self.assertTrue(any("fractional" in warning.lower() for warning in fractional.warnings))

        fractional.product_lines[0]["quantity"] = 5
        apply_delivery_docket_validation(fractional)
        self.assertTrue(fractional.importable)
        self.assertFalse(any("fractional" in warning.lower() for warning in fractional.warnings))


class DeliveryDocketRealCorpusTest(unittest.TestCase):
    def test_attached_production_format_corpus(self):
        cases = (
            (
                ["4376","DIRECT","SRGS PTY LTD","413 MT ATKINSON ROAD","TRUGANINA",None,[1680,450],6,0],
                (
                    "DELIVERY DOCKET: 4376/185531",
                    "DATED: 21/08/2026",
                    "DELIVER TO:",
                    "SRGS PTY LTD",
                    "MELBOURNE DISTRIBUTION CENTRE",
                    "STORE DC31",
                    "413 MT ATKINSON ROAD",
                    "TRUGANINA",
                    "Time slot:  FRIDAY 21/08/26  @ 9am",
                    "BOOKING # 2232671",
                    "ORDER NUMBER: 4522009179",
                    "1120 X 1.5KG COLOR T SHIRT RAGS",
                    "45 X 10KG COLOR T SHIRT RAGS",
                    "6 PALLETS",
                    "*INVOICE TO FOLLOW FROM Mcc ragman Pty ltd Trading as Melbourne cleaning cloths                PH:03 9930 7700",
                ),
            ),
            (
                ["4386","DIRECT","SRGS PTY LTD","413 MT ATKINSON ROAD","TRUGANINA",None,[1008,450,336],5,0],
                (
                    "DELIVERY DOCKET: 4386/185694",
                    "DATED: 28/08/2026",
                    "DELIVER TO:",
                    "SRGS PTY LTD",
                    "MELBOURNE DISTRIBUTION CENTRE",
                    "STORE DC31",
                    "413 MT ATKINSON ROAD",
                    "TRUGANINA",
                    "Time slot:  FRIDAY 28/08/26  @ 9am",
                    "BOOKING # 2233237",
                    "ORDER NUMBER: 4522039792",
                    "672 X 1.5KG COLOR T SHIRT RAGS",
                    "45 X 10KG COLOR T SHIRT RAGS",
                    "224 X WHITE COTTON RAGS 1.5KG",
                    "5 PALLETS",
                    "*INVOICE TO FOLLOW FROM Mcc ragman Pty ltd Trading as Melbourne cleaning cloths                PH:03 9930 7700",
                ),
            ),
            (
                ["4387","ON_FORWARD","RED BANTAM","36-38 GLENBARRY ROAD","CAMPBELLFIELD","0407 734 217",[800,400],3,0],
                (
                    "DELIVERY DOCKET: 4387/185705",
                    "DATED:  20/08/2026",
                    "EMAIL KEATING FREIGHT LINES TO ADVISE DROP OFF OF STOCK TO THEIR CAMPBELLFIELD DEPOT    office@keatingfreightlines.com.au",
                    "DELIVER to",
                    "Keating freight lines\t\t\tOPEN 7AM",
                    "36-38 glenbarry road",
                    "Campbellfield   9357 0206",
                    "ON FWD TO:",
                    "RED BANTAM",
                    "UNITS 4-5A/626 DALLINGER ROAD",
                    "LAVINGTON NSW 2641",
                    "Open mon-fri 9am-3pm",
                    "Call kate 0407 734 217     30mins prior delivery",
                    "STOCK:",
                    "40 x 20kg white BATH TOWEL",
                    "20 x 20kg colour bath towel",
                    "3 palletS",
                    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS  03 9930 7700",
                ),
            ),
            (
                ["4389","DIRECT","LIVLOR","8 TULLAMARINE PARK ROAD","TULLAMARINE",None,[50],0,5],
                (
                    "DELIVERY DOCKET: 4389/185763",
                    "DATED: 24/08/2026",
                    "DELIVER TO:",
                    "LIVLOR",
                    "8 TULLAMARINE PARK ROAD",
                    "TULLAMARINE",
                    "ORDER NUMBER: ALI",
                    "5 X 10KG PURE WHITE RAGS",
                ),
            ),
            (
                ["4391","ON_FORWARD","UNITED FASTENERS","13 SOMERS STREET","SUNSHINE NORTH","02 9131 3333",[630],1,0],
                (
                    "DELIVERY DOCKET: 4391/185816",
                    "DATED: 25/08/2026",
                    "VIA BURKINSHAWS TRANSPORT",
                    "EMAIL TRANSPORT NOTIFICATION OF DROP OFF",
                    "wagga@burkinshawstransport.com.au   (narelle)",
                    "DROP OFF TO:",
                    "C/-CHEMHAUL LOGISTICS",
                    "13 SOMERS STREET",
                    "SUNSHINE NORTH",
                    "8AM-3PM",
                    "ON FORWARD to:",
                    "UNITED FASTENERS",
                    "UNIT 2/3 BALL PL",
                    "WAGGA WAGGA    PH: 02 9131 3333",
                    "ORDER NUMBER: 40041845",
                    "63 X10kgs WORKSHOP RAGS",
                    "1 pallet",
                    "Please CHARGE  :MELBOURNE CLEANING CLOTHS ACCT#       PH: 03 9930 7700",
                ),
            ),
            (
                ["4392","ON_FORWARD","CAPITAL PAINT PLACE","47 ENDEAVOUR WAY","SUNSHINE","9313 9933",[440,190],1,0],
                (
                    "DELIVERY DOCKET: 4392/185817",
                    "DATED: 25/08/2026",
                    "EMAIL ABLETTS TO ADVISE DROP OFF OF STOCK TO THEIR SUNSHINE DEPOT    melbourneops@abletts.com.au   &  veronica@abletts.com.au",
                    "ABLETTS WILL SEND INVOICE ONCE PALLETS HAVE BEEN RECEIVED AT THEIR DEPOT FOR COD PMT",
                    "DELIVER to",
                    "ABLETTS TRANSPORT (TRANSPORT DEPOT)",
                    "47 ENDEAVOUR WAY",
                    "SUNSHINE VIC   ph: 9313 9933",
                    "ON FWD TO:",
                    "CAPITAL PAINT PLACE",
                    "18 ISA STREET",
                    "FYSHWICK ACT 2609",
                    "STOCK:",
                    "44 X 10KG  WHITE T SHIRT RAGS",
                    "19 X 10KG FLANNEL RAGS",
                    "1 PALLET",
                    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS PH: 03 9930 7700",
                ),
            ),
            (
                ["4393","ON_FORWARD","KONNECT FASTENING SYSTEMS","47 ENDEAVOUR WAY","SUNSHINE WEST","0428 723 989",[630],1,0],
                (
                    "DELIVERY DOCKET: 4393/18511",
                    "DATED: 25/08/2026",
                    "EMAIL SOUTH WEST FREIGHT TO ADVISE DROP OFF OF STOCK TO THEIR SUNSHINE DEPOT    operations@swfreight.com.au",
                    "DELIVER to",
                    "SOUTH WEST FREIGHT\t (SWF)\tOPEN 7AM",
                    "C/-CDM TRANSPORT",
                    "47 ENDEAVOUR WAY\t\t**please charge smiths rags acct**",
                    "SUNSHINE WEST VIC",
                    "ON FWD TO CUSTOMER:",
                    "Konnect fastening systems",
                    "321 commercial st west",
                    "Mount gambier sa 5290",
                    "Caine m: 0428 723 989",
                    "ORDER NUMBER:",
                    "63 x10kg workshop/mix cotton",
                    "1PALLET",
                    "*INVOICE TO FOLLOW FROM SMITHS rags/Melbourne cleaning cloths",
                    "PH: 03 9930 7730",
                ),
            ),
            (
                ["4394","DIRECT","SRGS PTY LTD","413 MT ATKINSON ROAD","TRUGANINA",None,[1344,1350],7,0],
                (
                    "DELIVERY DOCKET: 4394/185844",
                    "DATED: 04/09/2026",
                    "DELIVER TO:",
                    "SRGS PTY LTD",
                    "MELBOURNE DISTRIBUTION CENTRE",
                    "STORE DC31",
                    "413 MT ATKINSON ROAD",
                    "TRUGANINA",
                    "Time slot:  FRIDAY 04/09/26  @ 9am",
                    "BOOKING # 2233861",
                    "ORDER NUMBER: 4522061722",
                    "896 X 1.5KG COLOR T SHIRT RAGS",
                    "135 X 10KG COLOR T SHIRT RAGS",
                    "7 PALLETS",
                    "*INVOICE TO FOLLOW FROM Mcc ragman Pty ltd Trading as Melbourne cleaning cloths                PH:03 9930 7700",
                ),
            ),
            (
                ["4396","ON_FORWARD","SUPER CHEAP AUTO PTY LTD","GATE 1 /SHED 1, 19-43 ENTERPRIZE RD","WEST MELBOURNE",None,[336,450,336,450],5,0],
                (
                    "DELIVERY DOCKET: 4396/INVOICE#185858",
                    "DATED: 31/08/2026",
                    "DELIVER TO:",
                    "QUBE MELBOURNE\t\t\tbooking #qubm2608021263",
                    "GATE 1 /SHED 1, 19-43 ENTERPRIZE RD",
                    "(ENTRY OFF PITT ST) WEST MELBOURNE",
                    "*DRIVER MUST PRESENT COPY OF KUEHNE & NAGEL BOOKING CONFIRMATION **ATTACHED**",
                    "*GO TO GATE HOUSE & ADVISE  BOOKING HAS BEEN MADE* TO AVOID WAIT TIME",
                    "To: SUPER CHEAP AUTO PTY LTD",
                    "Auckland distribution centre  (dc04)",
                    "180 savill drive 2024 otahuhu",
                    "Aucklandbooking@superretailgroup.com",
                    "QUBE Time slot: MONDAY 31/08/2026 @ 10:00-10:59 am",
                    "K& N BOOKING# \t\tAU2608021263",
                    "K&n accounting #\t\t1076870568-0912",
                    "Tracking #\t\t\t1076 870 568",
                    "TOTAL: \t\t\t\t5 PALLETS",
                    "PURCHASE ORDER NUMBERS & PALLET BREAKDOWN",
                    "4522061448\t\t2 pallet\tCOLOUR RAGS RAGS 1.5KG\t           224 BAGS",
                    "4522061448\t\t1 pallet\tCOLOUR RAGS RAGS 10KG\t              45 BAGS",
                    "4522061448\t\t1 PALLET\tWHITE COTTON 1.5KG\t\t\t224 BAGS",
                    "4522061448\t\t1 PALLET\tWHITE T SHIRT 10KG\t\t\t45BAGS",
                    "From : Mcc ragman Pty ltd Trading as Melbourne cleaning cloths                            REGO # XW36ID DRIVER: NONDA TSATSOULIS",
                ),
            ),
            (
                ["4397","ON_FORWARD","STEAMATIC CANBERRA","47 ENDEAVOUR WAY","SUNSHINE","9313 9933",[450],1,0],
                (
                    "DELIVERY DOCKET: 4397/185869",
                    "DATED: 27/08/2026",
                    "EMAIL ABLETTS TO ADVISE DROP OFF OF STOCK TO THEIR SUNSHINE DEPOT    melbourneops@abletts.com.au   &  veronica@abletts.com.au",
                    "ABLETTS WILL SEND INVOICE ONCE PALLETS HAVE BEEN RECEIVED AT THEIR DEPOT FOR COD PMT",
                    "DELIVER to",
                    "ABLETTS TRANSPORT (TRANSPORT DEPOT)",
                    "47 ENDEAVOUR WAY",
                    "SUNSHINE VIC   ph: 9313 9933",
                    "ON FWD TO:",
                    "STEAMATIC CANBERRA",
                    "2/38 DACRE STREET",
                    "MITCHELL 2911",
                    "STOCK:",
                    "45 X 10KG  WHITE BATH TOWEL",
                    "1 PALLET",
                    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS PH: 03 9930 7700",
                ),
            ),
            (
                ["4398","DIRECT","REGENT RV","20-50 FILO DRIVE","SOMERTON","0477 511 802",[900],2,0],
                (
                    "DELIVERY DOCKET: 4398/185893",
                    "DATED: 27/08/2026",
                    "DELIVER TO:",
                    "REGENT RV",
                    "20-50 FILO DRIVE",
                    "SOMERTON",
                    "ATT: ATRA 0477 511 802",
                    "ORDER NUMBER: PO62211",
                    "90 X 10KG PURE WHITE RAGS",
                    "2 x PALLETS",
                    "*INVOICE TO FOLLOW FROM",
                    "MICRO FASTENERS",
                    "*CUSTOMER NAME: _______________SIGNATURE:_______________",
                    "DATE STOCK RECEIVED ______/______/2026",
                ),
            ),
            (
                ["4399","ON_FORWARD","KONNECT SHEPPARTON","36-38 GLENBARRY ROAD","CAMPBELLFIELD","9357 0206",[630],1,0],
                (
                    "DELIVERY DOCKET: 4399/185908",
                    "DATED: 28/08/2026",
                    "EMAIL KEATING FREIGHT LINES TO ADVISE DROP OFF OF STOCK TO THEIR CAMPBELLFIELD DEPOT    office@keatingfreightlines.com.au",
                    "DELIVER to",
                    "Keating freight lines\t\t\tOPEN 7AM",
                    "36-38 glenbarry road",
                    "Campbellfield   9357 0206",
                    "ON FWD TO:",
                    "KONNECT SHEPPARTON",
                    "84 BENALLA STREET",
                    "SHEPPARTON",
                    "STOCK:",
                    "63 X 10KG WORKSHOP rags",
                    "1 X PALLET",
                    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS",
                    "ph;03 99307700",
                ),
            ),
            (
                ["4400","ON_FORWARD","RED BANTAM","36-38 GLENBARRY ROAD","CAMPBELLFIELD","0407 734 217",[400],1,0],
                (
                    "DELIVERY DOCKET: 4400/185852",
                    "DATED:  28/08/2026",
                    "EMAIL KEATING FREIGHT LINES TO ADVISE DROP OFF OF STOCK TO THEIR CAMPBELLFIELD DEPOT    office@keatingfreightlines.com.au",
                    "DELIVER to",
                    "Keating freight lines\t\t\tOPEN 7AM",
                    "36-38 glenbarry road",
                    "Campbellfield   9357 0206",
                    "ON FWD TO:",
                    "RED BANTAM",
                    "UNITS 4-5A/626 DALLINGER ROAD",
                    "LAVINGTON NSW 2641",
                    "Open MON - THUR 9am-3pm",
                    "Call kate 0407 734 217     30mins prior delivery",
                    "STOCK:",
                    "20 x 20kg white BATH TOWEL",
                    "1 pallet",
                    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS  03 9930 7700",
                ),
            ),
            (
                ["4402","ON_FORWARD","KONNECT FASTENING SYSTEMS - WAGGA WAGGA","13 SOMERS STREET","SUNSHINE NORTH","02 6925 6700",[630],1,0],
                (
                    "DELIVERY DOCKET: 4402/185936",
                    "DATED: 31/08/2026",
                    "EMAIL BURKINSHAWS TO ADVISE DROP OFF OF STOCK TO THEIR SUNSHINE DEPOT    wagga@burkinshawstransport.com.au",
                    "DELIVER to",
                    "C/- CHEMHAUL LOGISTICS",
                    "13 SOMERS STREET",
                    "SUNSHINE NORTH",
                    "8AM-3PM",
                    "***MUST BE ON GOOD PALLET***",
                    "ON FWD TO:",
                    "KONNECT fastening systems - WAGGA WAGGA",
                    "49 PEARSON STREET          ** NEW ADDRESS**",
                    "WAGGA WAGGA 2650",
                    "PH: 02 6925 6700",
                    "STOCK:",
                    "63 X 10KG WORKSHOP RAGS",
                    "1 PALLET",
                    "INVOICE TO FOLLOW FROM SMITHS RAGS (MCC) PH: 03 9930 7700 (MCC)",
                ),
            ),
            (
                ["4404","ON_FORWARD","TOTAL TOOLS","36-38 GLENBARRY ROAD","CAMPBELLFIELD","5821 6400",[500,60],1,0],
                (
                    "DELIVERY DOCKET: 4404/186017",
                    "DATED: 02/09/2026",
                    "EMAIL KEATING FREIGHT LINES TO ADVISE DROP OFF OF STOCK TO THEIR CAMPBELLFIELD DEPOT    office@keatingfreightlines.com.au",
                    "DELIVER to",
                    "Keating freight lines\t\t\tOPEN 7AM",
                    "36-38 glenbarry road",
                    "Campbellfield   9357 0206",
                    "ON FWD TO:",
                    "TOTAL TOOLS",
                    "ATT: RAY  PH: 5821 6400",
                    "46-52 BENALLA RD",
                    "SHEPPARTON",
                    "STOCK:",
                    "50 X 10KG COLOURED RAGS",
                    "6 x 10kg windcheater",
                    "1 PALLET",
                    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS",
                    "ph;03 99307700",
                ),
            ),
            (
                ["4405","ON_FORWARD","AG PARTS","23 FOUNDATION ROAD","TRUGANINA","03 5482 4233",[900],2,0],
                (
                    "DELIVERY DOCKET: 4405/186015",
                    "DATED:  02/09/2026",
                    "DELIVER TO:",
                    "BURTCHELLS TSPT",
                    "C/- SARGEANT TRANSPORT",
                    "23 FOUNDATION ROAD",
                    "TRUGANINA",
                    "On forward to:",
                    "AG PARTS\t\t**PLEASE CHARGE CUSTOMERS ACCT**",
                    "39 MUNDARRA RD",
                    "ECHUCA VIC  3564    PH: 03 5482 4233",
                    "ORDER NUMBER: 40010686",
                    "90 X10KG COLOURED RAGS",
                    "2 PALLETS",
                    "INVOICE TO FOLLOW FROM mcc",
                ),
            ),
            (
                ["4408","ON_FORWARD","PIRTEK BENDIGO","36-38 GLENBARRY ROAD","CAMPBELLFIELD",None,[200],1,0],
                (
                    "DELIVERY DOCKET: 4408/186025",
                    "Email to book: office@keatingfreightlines.com.au",
                    "DATED: 02/09/2026",
                    "DELIVER to",
                    "keatings transport          OPEN 7AM- 3pm",
                    "36-38 glenbarry road",
                    "campbellfield",
                    "ON FWD TO:",
                    "PIRTEK BENDIGO",
                    "7B WELLSFORD DV",
                    "EAST BENDIGO",
                    "STOCK:",
                    "20 X 10KG WHITE BATH TOWEL RAGS",
                    "1 X PALLET",
                    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS 03 9930 7700",
                    "(please charge Melbourne cleaning cloths acct )",
                ),
            ),
            (
                ["4409","ON_FORWARD","K2 INDUSTRIAL SUPPLIES","BUILDING 3, 82 -86 BERKSHIRE RD","SUNSHINE NTH","5482 4446",[630],1,0],
                (
                    "DELIVERY DOCKET: 4409/186066",
                    "DATED: 03/09/2026",
                    "VIA SYMES TRANSPORT 5443 4199 RING TO ADVISE",
                    "DROP OFF TO SUNSHINE DEPOT BELOW",
                    "E: miCK@SYMESTRANSPORT.COM",
                    "DELIVERY TO:",
                    "SYMES TSPT C/- ECO WOOL INSULATION",
                    "BUILDING 3, 82 -86 BERKSHIRE RD SUNSHINE NTH",
                    "ENTER VIA STEERS ST THRU GATE ON RHS",
                    "ON FORWARD TO:    **PLEASE CHARGE CUSTOMER ACCT#",
                    "K2 INDUSTRIAL SUPPLIES",
                    "55 CORNELIA CREEK RD",
                    "ECHUCA  5482 4446",
                    "ORDER NUMBER: 30010822",
                    "63 X 10KG COLOUR RAGS",
                    "1 PALLET",
                    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS PH; 03 9930 7700",
                ),
            ),
            (
                ["4410","ON_FORWARD","RED BANTAM","36-38 GLENBARRY ROAD","CAMPBELLFIELD","0407 734 217",[800],2,0],
                (
                    "DELIVERY DOCKET: 4410/186073",
                    "DATED:  03/09/2026",
                    "EMAIL KEATING FREIGHT LINES TO ADVISE DROP OFF OF STOCK TO THEIR CAMPBELLFIELD DEPOT    office@keatingfreightlines.com.au",
                    "DELIVER to",
                    "Keating freight lines\t\t\tOPEN 7AM",
                    "36-38 glenbarry road",
                    "Campbellfield   9357 0206",
                    "ON FWD TO:",
                    "RED BANTAM",
                    "UNITS 4-5A/626 DALLINGER ROAD",
                    "LAVINGTON NSW 2641",
                    "Open MON - THUR 9am-3pm",
                    "Call kate 0407 734 217     30mins prior delivery",
                    "STOCK:",
                    "40 x 20kg white BATH TOWEL",
                    "2 pallets",
                    "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS  03 9930 7700",
                ),
            ),
        )
        for expected, lines in cases:
            with self.subTest(docket=expected[0]):
                parsed = _parse(lines, f"attached-{expected[0]}.docx")
                self.assertEqual(expected, [
                    parsed.docket_number, parsed.delivery_mode, parsed.company_name,
                    parsed.delivery_address, parsed.suburb, parsed.phone,
                    [line["quantity"] for line in parsed.product_lines],
                    parsed.pallet_quantity, parsed.loose_bags_quantity,
                ])
                self.assertIsNone(parsed.postcode)
                self.assertEqual(0, parsed.carton_quantity)
                self.assertEqual("2026-08-14", parsed.delivery_date)
                self.assertTrue(parsed.importable)
                if parsed.docket_number == "4396":
                    self.assertEqual("185858", parsed.invoice_number)
                    self.assertEqual("4522061448", parsed.order_no)
                    self.assertEqual(("10:00", "10:59"), (parsed.start_time, parsed.end_time))
                    self.assertIn("Deliver To: QUBE MELBOURNE\n", parsed.note)
                    self.assertIn("ENTRY OFF PITT ST", parsed.note)
                    self.assertIn("180 SAVILL DRIVE 2024 OTAHUHU", parsed.note)
                    self.assertIn("Booking #: qubm2608021263", parsed.note)

    def _parse(self, reference, physical, final=(), details=(), final_heading="ON FORWARD TO:"):
        return _parse([
            f"DELIVERY DOCKET: {reference}", "DATED: 28/08/2026", "DELIVERY TO:",
            *physical, *([final_heading, *final] if final else []), *details,
        ], f"corpus-{reference.split('/')[0]}.docx")

    def test_corpus_4376_standard_direct_control(self):
        parsed = self._parse("4376/185531", [
            "SRGS PTY LTD", "413 MT ATKINSON ROAD", "TRUGANINA",
        ], details=[
            "ORDER NUMBER: 4522009179", "1120 X 1.5KG COLOR T SHIRT RAGS",
            "45 X 10KG COLOR T SHIRT RAGS", "6 PALLETS",
        ])
        self.assertEqual(("DIRECT", "SRGS PTY LTD", "413 MT ATKINSON ROAD", "TRUGANINA"), (
            parsed.delivery_mode, parsed.company_name, parsed.delivery_address, parsed.suburb,
        ))
        self.assertEqual("4522009179", parsed.order_no)
        self.assertEqual(6, parsed.pallet_quantity)
        self.assertEqual([1680, 450], [line["quantity"] for line in parsed.product_lines])
        self.assertEqual(["COLOR T SHIRT RAGS"] * 2, [line["product_name"] for line in parsed.product_lines])
        self.assertTrue(parsed.importable)

    def test_corpus_4386_weight_at_end_preserves_all_three_products(self):
        parsed = self._parse("4386", [
            "SRGS PTY LTD", "413 MT ATKINSON ROAD", "TRUGANINA",
        ], details=[
            "672 X 1.5 KG COLOR T SHIRT RAGS", "45 X 10 KG COLOR T SHIRT RAGS",
            "224 X WHITE COTTON RAGS 1.5KG", "5 PALLETS",
        ])
        self.assertEqual([1008, 450, 336], [line["quantity"] for line in parsed.product_lines])
        self.assertEqual({
            "product_code": None, "product_name": "WHITE COTTON RAGS", "quantity": 336,
            "unit": "KG", "package_quantity": 224, "package_unit": "BAG1.5",
        }, parsed.product_lines[2])
        self.assertEqual(0, parsed.loose_bags_quantity)
        self.assertTrue(parsed.importable)

    def test_corpus_4387_final_contact_wins_over_inline_depot_phone(self):
        parsed = self._parse("4387", [
            "KEATING FREIGHT LINES", "36-38 GLENBARRY ROAD", "CAMPBELLFIELD 9357 0206",
        ], [
            "RED BANTAM", "UNITS 4-5A/626 DALLINGER ROAD", "LAVINGTON NSW 2641",
            "Call Kate 0407 734 217 30mins prior delivery",
        ], ["63 X 10KG COLOUR RAGS", "1 PALLET"])
        self.assertEqual("RED BANTAM", parsed.company_name)
        self.assertEqual("ON_FORWARD", parsed.delivery_mode)
        self.assertEqual("36-38 GLENBARRY ROAD", parsed.delivery_address)
        self.assertEqual("CAMPBELLFIELD", parsed.suburb)
        self.assertIsNone(parsed.postcode)
        self.assertEqual("0407 734 217", parsed.phone)
        self.assertIn("Contact: KATE", parsed.note)
        self.assertIn("30MINS PRIOR DELIVERY", parsed.note)
        self.assertIn("UNITS 4-5A/626 DALLINGER ROAD / LAVINGTON NSW 2641", parsed.note)
        self.assertTrue(parsed.importable)

    def test_corpus_4393_preserves_annotation_and_mobile_contact(self):
        parsed = self._parse("4393/18511", [
            "SOUTH WEST FREIGHT (SWF) OPEN 7AM", "C/-CDM TRANSPORT",
            "47 ENDEAVOUR WAY **please charge smiths rags acct**", "SUNSHINE WEST VIC",
        ], [
            "Konnect fastening systems", "321 commercial st west", "Mount gambier sa 5290",
            "Caine m: 0428 723 989",
        ], ["63 X 10KG WORKSHOP/MIX COTTON", "1PALLET"])
        self.assertEqual("KONNECT FASTENING SYSTEMS", parsed.company_name)
        self.assertEqual("47 ENDEAVOUR WAY", parsed.delivery_address)
        self.assertEqual("SUNSHINE WEST", parsed.suburb)
        self.assertEqual("0428 723 989", parsed.phone)
        self.assertIn("Contact: CAINE", parsed.note)
        self.assertIn("Annotation: PLEASE CHARGE SMITHS RAGS ACCT", parsed.note)
        self.assertIn("321 COMMERCIAL ST WEST / MOUNT GAMBIER SA 5290", parsed.note)
        self.assertTrue(parsed.importable)

    def test_corpus_4396_contextual_consignee_pallet_breakdown_and_time_range(self):
        parsed = self._parse("4396/INVOICE#185858", [
            "QUBE MELBOURNE", "GATE 1 /SHED 1, 19-43 ENTERPRIZE RD", "WEST MELBOURNE",
        ], [
            "Auckland distribution centre", "9 EXAMPLE ROAD", "AUCKLAND 2022",
        ], [
            "TOTAL: 5 PALLETS", "PURCHASE ORDER NUMBERS & PALLET BREAKDOWN",
            "4522061448 2 pallet COLOUR RAGS RAGS 1.5KG 224 BAGS",
            "4522061448 1 pallet COLOUR RAGS RAGS 10KG 45 BAGS",
            "4522061448 1 pallet WHITE COTTON 1.5KG 224 BAGS",
            "4522061448 1 pallet WHITE T SHIRT 10KG 45 BAGS",
            "QUBE Time slot: MONDAY 31/08/2026 @ 10:00-10:59 am",
        ], final_heading="To: SUPER CHEAP AUTO PTY LTD")
        self.assertEqual("4396", parsed.docket_number)
        self.assertEqual("INVOICE#185858", parsed.docket_reference)
        self.assertEqual("185858", parsed.invoice_number)
        self.assertEqual("ON_FORWARD", parsed.delivery_mode)
        self.assertEqual("SUPER CHEAP AUTO PTY LTD", parsed.company_name)
        self.assertEqual("GATE 1 /SHED 1, 19-43 ENTERPRIZE RD", parsed.delivery_address)
        self.assertEqual("WEST MELBOURNE", parsed.suburb)
        self.assertIsNone(parsed.postcode)
        self.assertEqual("4522061448", parsed.order_no)
        self.assertEqual(5, parsed.pallet_quantity)
        self.assertEqual(0, parsed.loose_bags_quantity)
        self.assertEqual([336, 450, 336, 450], [line["quantity"] for line in parsed.product_lines])
        self.assertEqual([224, 45, 224, 45], [line["package_quantity"] for line in parsed.product_lines])
        self.assertEqual(["BAG1.5", "BAG10", "BAG1.5", "BAG10"], [line["package_unit"] for line in parsed.product_lines])
        self.assertEqual(["COLOUR RAGS RAGS", "COLOUR RAGS RAGS", "WHITE COTTON", "WHITE T SHIRT"], [
            line["product_name"] for line in parsed.product_lines
        ])
        self.assertEqual(("10:00", "10:59"), (parsed.start_time, parsed.end_time))
        self.assertEqual("2026-08-14", parsed.delivery_date)
        self.assertIn("Deliver To: QUBE MELBOURNE", parsed.note)
        self.assertIn("AUCKLAND DISTRIBUTION CENTRE / 9 EXAMPLE ROAD / AUCKLAND 2022", parsed.note)
        self.assertTrue(parsed.importable)

    def test_corpus_4398_direct_attention_mobile(self):
        parsed = self._parse("4398", [
            "REGENT RV", "20-50 FILO DRIVE", "SOMERTON", "ATT: ATRA 0477 511 802",
        ], details=["1 PALLET"])
        self.assertEqual("REGENT RV", parsed.company_name)
        self.assertEqual("20-50 FILO DRIVE", parsed.delivery_address)
        self.assertEqual("SOMERTON", parsed.suburb)
        self.assertEqual("0477 511 802", parsed.phone)
        self.assertIn("Contact: ATRA", parsed.note)
        self.assertTrue(parsed.importable)

    def test_corpus_4402_remote_new_address_annotation_stays_in_note(self):
        parsed = self._parse("4402", [
            "C/- CHEMHAUL LOGISTICS", "13 SOMERS STREET", "SUNSHINE NORTH",
        ], [
            "KONNECT FASTENING SYSTEMS - WAGGA WAGGA",
            "49 PEARSON STREET ** NEW ADDRESS**", "WAGGA WAGGA NSW 2650",
        ], ["1 PALLET"])
        self.assertEqual("KONNECT FASTENING SYSTEMS - WAGGA WAGGA", parsed.company_name)
        self.assertEqual("13 SOMERS STREET", parsed.delivery_address)
        self.assertEqual("SUNSHINE NORTH", parsed.suburb)
        self.assertIn("On Forward Address: 49 PEARSON STREET / WAGGA WAGGA NSW 2650", parsed.note)
        self.assertIn("Annotation: NEW ADDRESS", parsed.note)
        self.assertTrue(parsed.importable)

    def test_corpus_4404_final_attention_and_phone(self):
        parsed = self._parse("4404", [
            "KEATING FREIGHT LINES", "36-38 GLENBARRY ROAD", "CAMPBELLFIELD",
        ], [
            "TOTAL TOOLS", "1 EXAMPLE ROAD", "SHEPPARTON VIC 3630", "ATT: RAY PH: 5821 6400",
        ], ["1 PALLET"])
        self.assertEqual("TOTAL TOOLS", parsed.company_name)
        self.assertEqual("36-38 GLENBARRY ROAD", parsed.delivery_address)
        self.assertEqual("CAMPBELLFIELD", parsed.suburb)
        self.assertEqual("5821 6400", parsed.phone)
        self.assertIn("Contact: RAY", parsed.note)
        self.assertNotIn("On Forward Address: ATT", parsed.note)
        self.assertTrue(parsed.importable)

    def test_corpus_4405_long_forward_annotated_company_and_inline_postcode_phone(self):
        parsed = self._parse("4405", [
            "BURTCHELLS TSPT", "C/- SARGEANT TRANSPORT", "23 FOUNDATION ROAD", "TRUGANINA",
        ], [
            "AG PARTS **PLEASE CHARGE CUSTOMERS ACCT**", "1 EXAMPLE ROAD",
            "ECHUCA VIC 3564 PH: 03 5482 4233",
        ], ["1 PALLET"], final_heading="ON FORWARD TO CUSTOMER:")
        self.assertEqual("AG PARTS", parsed.company_name)
        self.assertEqual("23 FOUNDATION ROAD", parsed.delivery_address)
        self.assertEqual("TRUGANINA", parsed.suburb)
        self.assertIsNone(parsed.postcode)
        self.assertEqual("03 5482 4233", parsed.phone)
        self.assertIn("C/- SARGEANT TRANSPORT", parsed.note)
        self.assertIn("Annotation: PLEASE CHARGE CUSTOMERS ACCT", parsed.note)
        self.assertIn("1 EXAMPLE ROAD / ECHUCA VIC 3564", parsed.note)
        self.assertTrue(parsed.importable)

    def test_corpus_4408_dv_suffix_and_final_address_context(self):
        parsed = self._parse("4408", [
            "KEATING FREIGHT LINES", "36-38 GLENBARRY ROAD", "CAMPBELLFIELD",
        ], ["PIRTEK BENDIGO", "7B WELLSFORD DV", "BENDIGO VIC 3550"], ["1 PALLET"])
        self.assertEqual("PIRTEK BENDIGO", parsed.company_name)
        self.assertEqual("36-38 GLENBARRY ROAD", parsed.delivery_address)
        self.assertEqual("CAMPBELLFIELD", parsed.suburb)
        self.assertIn("On Forward Address: 7B WELLSFORD DV / BENDIGO VIC 3550", parsed.note)
        self.assertTrue(parsed.importable)

    def test_corpus_4412_simple_direct_control(self):
        parsed = self._parse("4412", [
            "BEYOND PLUMBING SUPPLIES", "UNIT 1/14 TASMAN CT", "KEYSBOROUGH",
        ], details=["ORDER NUMBER: ALI", "1 PALLET"])
        self.assertEqual("DIRECT", parsed.delivery_mode)
        self.assertEqual("BEYOND PLUMBING SUPPLIES", parsed.company_name)
        self.assertEqual("UNIT 1/14 TASMAN CT", parsed.delivery_address)
        self.assertEqual("KEYSBOROUGH", parsed.suburb)
        self.assertEqual("ALI", parsed.order_no)
        self.assertTrue(parsed.importable)

    def test_contextual_to_requires_physical_and_final_address_context(self):
        for physical, final in (
            ([], ["REMOTE CUSTOMER", "1 EXAMPLE ROAD", "AUCKLAND 2022"]),
            (["DEPOT", "WEST MELBOURNE"], ["REMOTE CUSTOMER", "1 EXAMPLE ROAD", "AUCKLAND 2022"]),
            (["DEPOT", "19 ENTERPRIZE RD", "WEST MELBOURNE"], ["PLEASE CALL FIRST"]),
            (["DEPOT", "19 ENTERPRIZE RD", "WEST MELBOURNE"], ["UNRECOGNIZED ADDRESS FORMAT"]),
        ):
            with self.subTest(physical=physical, final=final):
                parsed = self._parse("4990", physical, final, ["1 PALLET"], final_heading="To: FINAL CUSTOMER")
                self.assertFalse(parsed.importable)
                self.assertIsNone(parsed.company_name)
                self.assertNotEqual("1 EXAMPLE ROAD", parsed.delivery_address)

    def test_email_instruction_and_footer_to_are_not_final_customers(self):
        for extra in (
            ["To: bookings@example.test", "PLEASE CALL FIRST"],
            ["To: PLEASE CALL DISPATCH", "1 EXAMPLE ROAD", "AUCKLAND 2022"],
            ["INVOICE TO FOLLOW FROM MCC", "To: ACCOUNTS CUSTOMER", "1 EXAMPLE ROAD", "AUCKLAND 2022"],
        ):
            with self.subTest(extra=extra):
                parsed = self._parse("4990", ["DEPOT", "19 ENTERPRIZE RD", "WEST MELBOURNE"], details=["1 PALLET", *extra])
                self.assertEqual("DIRECT", parsed.delivery_mode)
                self.assertEqual("DEPOT", parsed.company_name)
                self.assertEqual("19 ENTERPRIZE RD", parsed.delivery_address)
                self.assertTrue(parsed.importable)

    def test_empty_explicit_final_section_does_not_fall_back_to_depot_company(self):
        parsed = self._parse("4990", ["DEPOT", "19 ENTERPRIZE RD", "WEST MELBOURNE"],
                             ["PLEASE CALL DISPATCH", "1 EXAMPLE ROAD", "AUCKLAND 2022"], ["1 PALLET"])
        self.assertIsNone(parsed.company_name)
        self.assertFalse(parsed.importable)

    def test_empty_contextual_to_does_not_infer_company_from_locality(self):
        parsed = self._parse("4990", ["QUBE MELBOURNE", "19 EXAMPLE ROAD", "WEST MELBOURNE"],
                             ["AUCKLAND 2022"], ["1 PALLET"], final_heading="To:")
        self.assertIsNone(parsed.company_name)
        self.assertEqual("19 EXAMPLE ROAD", parsed.delivery_address)
        self.assertEqual("WEST MELBOURNE", parsed.suburb)
        self.assertFalse(parsed.importable)

    def test_instruction_and_contact_lines_never_become_locations(self):
        for instruction in (
            "DRIVER MUST USE SIDE ST", "GO TO THE REAR GATE", "MUST CALL FIRST", "PLEASE CALL FIRST",
            "CALL KATE 0407 734 217 30MINS PRIOR DELIVERY", "OPEN MON-FRI 7AM-3PM",
            "OPEN MON-THUR 7AM-3PM", "30MINS PRIOR DELIVERY", "ATT: ATRA 0477 511 802",
            "ATT: RAY PH: 5821 6400", "CAINE M: 0428 723 989",
        ):
            with self.subTest(instruction=instruction):
                parsed = self._parse("4990", [instruction, "TEST CUSTOMER", instruction,
                    "7B WELLSFORD DV", instruction, "BENDIGO VIC 3550"], details=["1 PALLET"])
                self.assertEqual("TEST CUSTOMER", parsed.company_name)
                self.assertEqual("7B WELLSFORD DV", parsed.delivery_address)
                self.assertEqual("BENDIGO", parsed.suburb)
                self.assertEqual("3550", parsed.postcode)
                self.assertTrue(parsed.importable)

    def test_phone_labels_postcodes_and_footer_precedence(self):
        for line, suburb, postcode, phone in (
            ("ECHUCA PH; 03 5482 4233", "ECHUCA", None, "03 5482 4233"),
            ("ECHUCA VIC 3564 PH: 03 5482 4233", "ECHUCA", "3564", "03 5482 4233"),
            ("CAMPBELLFIELD 9357 0206", "CAMPBELLFIELD", None, "9357 0206"),
            ("ALTONA VIC 3018", "ALTONA", "3018", None),
            ("BUNBURY WA 6230", "BUNBURY", "6230", None),
            ("ECHUCA TELEPHONE: 03 5482 4233", "ECHUCA", None, "03 5482 4233"),
        ):
            with self.subTest(line=line):
                parsed = self._parse("4990", ["TEST CUSTOMER", "1 EXAMPLE ROAD", line],
                                     details=["1 PALLET", "SMITHS RAGS PH; 03 9930 7700"])
                self.assertEqual((suburb, postcode, phone), (parsed.suburb, parsed.postcode, parsed.phone))
                self.assertNotIn("03 9930 7700", parsed.note)
        for footer in ("SMITHS RAGS PH: 03 9930 7700", "MCC RAGMAN PTY LTD PH: 03 9930 7700",
                       "INVOICE TO FOLLOW FROM MELBOURNE CLEANING CLOTHS PH; 03 9930 7700"):
            with self.subTest(footer=footer):
                parsed = self._parse("4990", ["TEST CUSTOMER", "1 EXAMPLE ROAD", "RICHMOND", footer],
                                     details=["1 PALLET"])
                self.assertIsNone(parsed.phone)
                self.assertNotIn("03 9930 7700", parsed.note)

    def test_pallet_breakdown_requires_structure_and_preserves_multiple_orders(self):
        physical = ["TEST CUSTOMER", "1 EXAMPLE ROAD", "RICHMOND"]
        rows = ["4522061448 2 pallet COLOUR RAGS RAGS 1.5KG 224 BAGS",
                "4522061449 1 pallet WHITE T SHIRT 10KG 45 BAGS"]
        for summary, expected in (([], 3), (["3 PALLETS"], 3), (["TOTAL: 5 PALLETS"], 5), (["TOTAL:", "5 PALLETS"], 5)):
            with self.subTest(summary=summary):
                parsed = self._parse("4990", physical, details=[*summary, "PURCHASE ORDER NUMBERS & PALLET BREAKDOWN", *rows])
                self.assertEqual(expected, parsed.pallet_quantity)
                self.assertEqual([336, 450], [line["quantity"] for line in parsed.product_lines])
                self.assertIsNone(parsed.order_no)
                self.assertIn("Purchase Orders: 4522061448, 4522061449", parsed.note)
        ordinary = self._parse("4990", physical, details=["ORDER NUMBER: EXPLICIT/PO",
            "PURCHASE ORDER NUMBERS & PALLET BREAKDOWN", rows[0]])
        self.assertEqual("EXPLICIT/PO", ordinary.order_no)
        for malformed in (
            "2 pallet COLOUR RAGS 1.5KG 224 BAGS", "4522061448 COLOUR RAGS 1.5KG 224 BAGS",
            "4522061448 2 pallet 1.5KG 224 BAGS", "4522061448 2 pallet COLOUR RAGS 224 BAGS",
            "4522061448 2 pallet COLOUR RAGS 1.5KG", "4522061448 2 pallet COLOUR RAGS 1.5KG 224",
        ):
            with self.subTest(malformed=malformed):
                parsed = self._parse("4990", physical, details=["PURCHASE ORDER NUMBERS & PALLET BREAKDOWN", malformed])
                self.assertEqual([], parsed.product_lines)
                self.assertEqual(0, parsed.pallet_quantity)
                self.assertIsNone(parsed.order_no)
                self.assertFalse(parsed.importable)
        no_heading = self._parse("4990", physical, details=rows)
        self.assertEqual([], no_heading.product_lines)
        self.assertIsNone(no_heading.order_no)

    def test_invoice_reference_and_load_variants(self):
        for reference, invoice in (
            ("186066", "186066"), ("INVOICE#185858", "185858"), ("invoice #185858", "185858"),
            ("INVOICE#18585", None), ("ORDER#185858", None), ("NEWWAY 185858", None), ("NEWAY 182-2", None),
        ):
            with self.subTest(reference=reference):
                parsed = self._parse(f"4990/{reference}", ["TEST CUSTOMER", "1 EXAMPLE ROAD", "RICHMOND"], details=["1 PALLET"])
                self.assertEqual(reference, parsed.docket_reference)
                self.assertEqual(invoice, parsed.invoice_number)
        for load, expected in (("1 PALLET", 1), ("1pallet", 1), ("1 X PALLET", 1),
                               ("2 PALLETS", 2), ("3 palletS", 3), ("TOTAL: 5 PALLETS", 5)):
            with self.subTest(load=load):
                parsed = self._parse("4990", ["TEST CUSTOMER", "1 EXAMPLE ROAD", "RICHMOND"], details=[load])
                self.assertEqual(expected, parsed.pallet_quantity)

    def test_time_slot_prefix_range_and_shared_meridiem(self):
        for slot, expected in (
            ("Time slot: FRIDAY 28/08/26 @ 9am", ("09:00", None)),
            ("QUBE Time slot: MONDAY 31/08/2026 @ 10:00-10:59 am", ("10:00", "10:59")),
            ("CARRIER Time slot: MONDAY 31/08/2026 @ 1:00-1:59 pm", ("13:00", "13:59")),
            ("Time slot: MONDAY 31/08/2026 @ 11am-1pm", ("11:00", "13:00")),
        ):
            with self.subTest(slot=slot):
                parsed = self._parse("4990", ["TEST CUSTOMER", "1 EXAMPLE ROAD", "RICHMOND"], details=[slot, "1 PALLET"])
                self.assertEqual(expected, (parsed.start_time, parsed.end_time))
                self.assertEqual("2026-08-14", parsed.delivery_date)


if __name__ == "__main__":
    unittest.main()
